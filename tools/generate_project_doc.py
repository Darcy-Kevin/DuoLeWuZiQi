#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成项目介绍Word文档
用于会议分享
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime


def create_project_document():
    """创建项目介绍文档"""
    doc = Document()

    # 设置文档默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)

    # ========== 封面 ==========
    # 标题
    title = doc.add_heading("多乐五子棋 UI 自动化测试项目", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.bold = True

    # 副标题
    subtitle = doc.add_paragraph("项目介绍与成果分享")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ========== 目录 ==========
    doc.add_heading("目录", 1)
    doc.add_paragraph("1. 项目概述", style="List Number")
    doc.add_paragraph("2. 技术架构", style="List Number")
    doc.add_paragraph("3. 核心功能", style="List Number")
    doc.add_paragraph("4. 测试覆盖", style="List Number")
    doc.add_paragraph("5. 项目成果", style="List Number")
    doc.add_paragraph("6. 发现的问题", style="List Number")
    doc.add_paragraph("7. 后续规划", style="List Number")

    doc.add_page_break()

    # ========== 1. 项目概述 ==========
    doc.add_heading("1. 项目概述", 1)

    para1 = doc.add_paragraph()
    para1.add_run("本项目是针对多乐五子棋应用（包名：").font.size = Pt(11)
    para1.add_run("com.duole.wuziqihd").font.bold = True
    para1.add_run(
        "）的移动端 UI 自动化测试框架，旨在通过自动化测试提升测试效率，保障应用质量，减少人工回归测试成本。"
    )

    doc.add_paragraph("项目目标：", style="List Bullet")
    doc.add_paragraph("• 实现关键业务流程的自动化回归测试", style="List Bullet 2")
    doc.add_paragraph("• 提升测试执行效率，缩短测试周期", style="List Bullet 2")
    doc.add_paragraph(
        "• 通过自动化测试发现潜在缺陷，保障产品质量", style="List Bullet 2"
    )
    doc.add_paragraph("• 建立可维护、可扩展的自动化测试体系", style="List Bullet 2")

    # ========== 2. 技术架构 ==========
    doc.add_heading("2. 技术架构", 1)

    doc.add_heading("2.1 技术栈", 2)

    tech_table = doc.add_table(rows=6, cols=2)
    tech_table.style = "Light Grid Accent 1"

    tech_data = [
        ("测试框架", "pytest 8.4.2"),
        ("UI自动化", "uiautomator2 3.4.2"),
        ("图像识别", "OpenCV 4.12.0.88"),
        ("测试报告", "Allure 2.15.0"),
        ("HTTP请求", "requests 2.32.5, httpx 0.28.1"),
        ("开发语言", "Python 3.9+"),
    ]

    for i, (category, tech) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = category
        tech_table.rows[i].cells[1].text = tech

    doc.add_heading("2.2 架构设计", 2)

    doc.add_paragraph("项目采用分层架构设计，主要包含以下模块：", style="List Bullet")
    doc.add_paragraph("• ", style="List Bullet 2")
    doc.add_paragraph(
        "  测试用例层（src/tests/）：包含各业务模块的测试用例", style="List Bullet 3"
    )
    doc.add_paragraph("• ", style="List Bullet 2")
    doc.add_paragraph(
        "  工具层（src/utils/）：提供截图、图像匹配、滑动等通用能力",
        style="List Bullet 3",
    )
    doc.add_paragraph("• ", style="List Bullet 2")
    doc.add_paragraph(
        "  配置层（src/config/）：统一管理坐标、包名等配置信息", style="List Bullet 3"
    )
    doc.add_paragraph("• ", style="List Bullet 2")
    doc.add_paragraph(
        "  资源层（src/resources/）：存储图像模板、测试截图等资源",
        style="List Bullet 3",
    )
    doc.add_paragraph("• ", style="List Bullet 2")
    doc.add_paragraph(
        "  脚本层（scripts/）：提供环境配置、测试执行、报告管理等脚本",
        style="List Bullet 3",
    )

    # ========== 3. 核心功能 ==========
    doc.add_heading("3. 核心功能", 1)

    doc.add_heading("3.1 设备驱动自动管理", 2)
    doc.add_paragraph(
        "通过 uiautomator2 框架实现 Android 设备的自动连接与初始化，支持真机和模拟器，脚本自动检测可用设备并建立连接。"
    )

    doc.add_heading("3.2 多场景图像识别", 2)
    doc.add_paragraph(
        "基于 OpenCV 模板匹配技术，实现按钮、红点、标题等界面元素的精准定位，支持自定义模板生成工具，便于扩展识别覆盖面。"
    )

    doc.add_heading("3.3 API + UI 联动验证", 2)
    doc.add_paragraph(
        "在邮件红点等场景中，通过接口造数触发业务逻辑，再通过 UI 自动化验证界面表现，实现端到端的业务闭环验证。"
    )

    doc.add_heading("3.4 Allure 报告集成", 2)
    doc.add_paragraph(
        "测试执行后自动生成结构化的 Allure 报告，包含测试步骤、截图、标记图等详细信息，便于问题定位和测试结果追溯。"
    )

    doc.add_heading("3.5 脚本化流程", 2)
    doc.add_paragraph(
        "提供一键环境配置、批量/单例运行、截图清理等辅助脚本，降低使用门槛，提升测试执行效率。"
    )

    # ========== 4. 测试覆盖 ==========
    doc.add_heading("4. 测试覆盖", 1)

    doc.add_heading("4.1 测试场景", 2)

    scenario_table = doc.add_table(rows=5, cols=2)
    scenario_table.style = "Light Grid Accent 1"

    scenario_data = [
        ("测试模块", "测试场景"),
        ("应用启动", "验证应用启动成功、包名匹配"),
        ("邮件功能", "邮件红点展示、邮件详情查看、红点消失验证"),
        ("每日签到", "签到入口红点、签到按钮点击、签到弹窗验证"),
        ("应用审核", "应用审核流程相关测试场景"),
    ]

    for i, (module, scenario) in enumerate(scenario_data):
        scenario_table.rows[i].cells[0].text = module
        scenario_table.rows[i].cells[1].text = scenario

    doc.add_heading("4.2 测试用例统计", 2)
    doc.add_paragraph("• 测试模块数量：4 个（应用启动、邮件、签到、审核）")
    doc.add_paragraph("• 测试用例总数：29+ 个")
    doc.add_paragraph("• 图像模板数量：89+ 个")

    # ========== 5. 项目成果 ==========
    doc.add_heading("5. 项目成果", 1)

    doc.add_paragraph("通过自动化测试的实施，项目取得了以下成果：", style="List Bullet")
    doc.add_paragraph(
        "• 提升了测试执行效率，自动化测试可在短时间内完成回归验证",
        style="List Bullet 2",
    )
    doc.add_paragraph(
        "• 建立了可维护的测试框架，便于后续扩展和维护", style="List Bullet 2"
    )
    doc.add_paragraph(
        "• 实现了关键业务流程的自动化覆盖，减少人工测试成本", style="List Bullet 2"
    )
    doc.add_paragraph(
        "• 通过自动化测试发现并记录了多个问题，提升了产品质量", style="List Bullet 2"
    )

    # ========== 6. 发现的问题 ==========
    doc.add_heading("6. 发现的问题", 1)

    doc.add_paragraph("在自动化测试执行过程中，共发现以下问题：", style="List Bullet")

    # 创建一个表格用于填写问题统计
    issue_table = doc.add_table(rows=2, cols=4)
    issue_table.style = "Light Grid Accent 1"

    # 表头
    header_cells = issue_table.rows[0].cells
    header_cells[0].text = "问题类型"
    header_cells[1].text = "严重程度"
    header_cells[2].text = "问题数量"
    header_cells[3].text = "备注"

    # 设置表头样式
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # 示例行（用户可以修改）
    data_cells = issue_table.rows[1].cells
    data_cells[0].text = "功能缺陷"
    data_cells[1].text = "高/中/低"
    data_cells[2].text = "待填写"
    data_cells[3].text = "待填写"

    doc.add_paragraph()
    doc.add_paragraph("问题详细列表：", style="List Bullet")

    # 添加一个空表格用于填写详细问题
    detail_table = doc.add_table(rows=2, cols=5)
    detail_table.style = "Light Grid Accent 1"

    detail_header = detail_table.rows[0].cells
    detail_header[0].text = "序号"
    detail_header[1].text = "问题描述"
    detail_header[2].text = "严重程度"
    detail_header[3].text = "发现时间"
    detail_header[4].text = "状态"

    for cell in detail_header:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # 示例行
    detail_data = detail_table.rows[1].cells
    detail_data[0].text = "1"
    detail_data[1].text = "（请在此填写问题描述）"
    detail_data[2].text = "（高/中/低）"
    detail_data[3].text = "（发现日期）"
    detail_data[4].text = "（已修复/待修复）"

    doc.add_paragraph()
    doc.add_paragraph("说明：", style="List Bullet")
    doc.add_paragraph(
        "• 请在上述表格中填写自动化测试发现的具体问题", style="List Bullet 2"
    )
    doc.add_paragraph(
        "• 问题类型可包括：功能缺陷、UI异常、性能问题、兼容性问题等",
        style="List Bullet 2",
    )
    doc.add_paragraph("• 严重程度分为：高、中、低三个级别", style="List Bullet 2")
    doc.add_paragraph("• 建议附上 Allure 报告中的截图和日志信息", style="List Bullet 2")

    # ========== 7. 后续规划 ==========
    doc.add_heading("7. 后续规划", 1)

    doc.add_paragraph("为了持续提升自动化测试能力，后续计划包括：", style="List Bullet")
    doc.add_paragraph(
        "• 扩展测试覆盖范围，增加更多业务场景的自动化测试", style="List Bullet 2"
    )
    doc.add_paragraph("• 优化测试执行效率，减少测试执行时间", style="List Bullet 2")
    doc.add_paragraph("• 完善测试报告，增加更多维度的数据分析", style="List Bullet 2")
    doc.add_paragraph(
        "• 建立持续集成流程，实现自动化测试的定时执行", style="List Bullet 2"
    )
    doc.add_paragraph(
        "• 提升测试稳定性，减少因环境因素导致的测试失败", style="List Bullet 2"
    )

    # ========== 附录 ==========
    doc.add_page_break()
    doc.add_heading("附录", 1)

    doc.add_heading("A. 项目结构", 2)
    doc.add_paragraph("scripts/run/          - 运行用例脚本")
    doc.add_paragraph("scripts/report/       - 报告管理脚本")
    doc.add_paragraph("src/tests/           - 测试用例")
    doc.add_paragraph("src/utils/           - 工具库")
    doc.add_paragraph("src/config/          - 配置文件")
    doc.add_paragraph("src/resources/       - 资源文件")

    doc.add_heading("B. 使用说明", 2)
    doc.add_paragraph("1. 环境初始化：./scripts/setup.sh")
    doc.add_paragraph("2. 运行全量测试：./scripts/run/run_tests.sh")
    doc.add_paragraph("3. 运行单用例：./scripts/run/run_single_test.sh <用例路径>")
    doc.add_paragraph("4. 查看报告：allure serve allure-results")

    # 保存文档
    output_path = "artifacts/项目介绍文档.docx"
    doc.save(output_path)
    print(f"✅ 文档已生成：{output_path}")
    print('📝 请在文档的"6. 发现的问题"章节中填写具体的问题信息')


if __name__ == "__main__":
    try:
        create_project_document()
    except ImportError:
        print("❌ 缺少 python-docx 库，正在安装...")
        import subprocess
        import sys

        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("✅ 安装完成，重新运行脚本...")
        create_project_document()
